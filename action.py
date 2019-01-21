from __future__ import print_function
from concurrent.futures import ThreadPoolExecutor, wait
from threading import Lock

from docx import Document
from watson_developer_cloud import NaturalLanguageUnderstandingV1
from watson_developer_cloud.natural_language_understanding_v1 import Features, ConceptsOptions, KeywordsOptions

import src.bsk_utils as utils
import json


class Action(object):
    """Represents an action that is to be taken on some sort of Box object"""

    def __init__(self, config: dict) -> 'Action':
        """
          Input:
            * Config dict
          Output:

          This function is the constructor of the Action that will store the config and model
        """

        self.config = config
        self.cards = []
        self.concepts = []
        self.keywords = []
        self.lock = Lock()
        self.model = NaturalLanguageUnderstandingV1(
            version='2018-03-16',
            iam_apikey=config['nlu_iam_key'],
            url=config['url'])

    def do_action(self, file_path: str, identification: int) -> None:
        """
          Input:
            * Path to file to run action on.
            * File Identification to append to the Box skill.
          Output:
            * Path to new file if exists. (Optional)

          This function is to call your Watson API on the input file
          and store the insights.

          In this function, add all of the cards to the self.cards list, these will
          later be added to the files that call this custom skill.
          The cards can be created in the correct format using the bsk_utils functions.
        """

        # Try to get the file as a docx file, if that fails the file cannot be processed here and a blank response is returned
        try:
            original_doc = Document(file_path)

            # Store all text lines in an array
            texts = []
            for table in original_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text != "":
                            texts.append(cell.text)

            for paragraph in original_doc.paragraphs:
                if paragraph.text != "":
                    runs = paragraph.runs
                    for run in runs:
                        if run.text != "":
                            texts.append(run.text)

            texts = '. '.join(texts)
        except ValueError as e:
            print("ERROR: Requested file {} is not a .docx file and cannot be processed".format(file_path))
            return
        text_collection = []

        self._parallel_NlU(texts)

        # Create the cards and add them to the list of cards to append to the file
        if self.concepts:
            concept_card = utils.create_keyword_card(
                self.concepts, "Concepts", identification, 1)
            self.cards.append(concept_card)

        if self.keywords:
            keyword_card = utils.create_keyword_card(
                self.keywords, "Keywords", identification, 1)
            self.cards.append(keyword_card)

    def push2storage(self, storage: 'Storage') -> list:
        """
          Input:
            * Storage Object
            * File Data Dict of original file in the format: {name: <FILENAME>, id: <FILE ID>}
          Output:
            * File name (updated file)

          This function updates a file to storage using the insights
          gained from the Watson API.
        """

        # Update the metadata of the file for each Box Skillcard, return the results for
        # debugging

        return storage.update_file({'cards': self.cards})

    def _parallel_NlU(self, text):
            
            # A Function to call Watson Natural Language Understanding

        if self.config['keywords']:
            keyword_option = KeywordsOptions(limit=self.config['keyword_limit'])
        else:
            keyword_option = None

        if self.config['concepts']:
            concepts_option = ConceptsOptions(
                limit=self.config['concept_limit'])
        else:
            concepts_option = None

        try:
            results = self.model.analyze(
                text=text,
                features=Features(
                    concepts=concepts_option,
                    keywords=keyword_option),
                language='en'
            )

            json_results = results.get_result()

            our_concepts = []
            for concept in json_results['concepts']:
                our_concepts.append(concept['text'])
            
            our_keywords = []
            for keyword in json_results['keywords']:
                our_keywords.append(keyword['text'])
            
            self.lock.acquire()
            self.concepts = self.concepts + our_concepts
            self.keywords = self.keywords + our_keywords
            self.lock.release()
        
        except Exception as e:
            print(str(e))
